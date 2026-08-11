import json
from pathlib import Path

import pytest
from docx import Document

from medical_checklist_verifier.checklist_digester import digest_checklist
from medical_checklist_verifier.checklist_digester.cli import main, validate_digest


def make_checklist(path: Path) -> None:
    document = Document()
    document.add_heading("Submission Checklist", level=1)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Requirement"
    table.cell(1, 0).text = "1. Identity"
    table.cell(1, 1).text = "Provide the legal manufacturer name."
    document.save(str(path))


def test_cli_writes_agent_ready_json_to_requested_path(tmp_path: Path) -> None:
    source = tmp_path / "input.docx"
    output = tmp_path / "digest.json"
    make_checklist(source)

    exit_code = main([str(source), "--output", str(output)])

    assert exit_code == 0
    payload = json.loads(output.read_text(encoding="utf-8"))
    assert payload["statistics"] == {
        "item_count": 1,
        "requirement_count": 1,
        "explicit_requirement_count": 1,
        "normative_candidate_count": 0,
        "guidance_count": 0,
    }
    assert payload["items"][0]["requirements"][0]["text"] == (
        "Provide the legal manufacturer name."
    )


def test_cli_prints_json_when_output_is_omitted(tmp_path: Path, capsys) -> None:
    source = tmp_path / "input.docx"
    make_checklist(source)

    exit_code = main([str(source)])

    assert exit_code == 0
    payload = json.loads(capsys.readouterr().out)
    assert payload["checklist"]["title"] == "Submission Checklist"


def test_cli_rejects_document_without_detectable_items(tmp_path: Path, capsys) -> None:
    source = tmp_path / "empty.docx"
    document = Document()
    document.add_paragraph("This is not a checklist.")
    document.save(str(source))

    exit_code = main([str(source)])

    assert exit_code == 2
    assert "No checklist items were detected" in capsys.readouterr().err


def test_cli_refuses_to_overwrite_source_document(tmp_path: Path, capsys) -> None:
    source = tmp_path / "input.docx"
    make_checklist(source)
    original = source.read_bytes()

    exit_code = main([str(source), "--output", str(source)])

    assert exit_code == 2
    assert "must differ from the input" in capsys.readouterr().err
    assert source.read_bytes() == original


def test_cli_fails_closed_on_unknown_checklist_table_layout(
    tmp_path: Path, capsys
) -> None:
    source = tmp_path / "unknown.docx"
    output = tmp_path / "unknown.json"
    document = Document()
    document.add_heading("Unknown Checklist", level=1)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Alpha"
    table.cell(0, 1).text = "Beta"
    table.cell(1, 0).text = "One"
    table.cell(1, 1).text = "Two"
    document.save(str(source))

    exit_code = main([str(source), "--output", str(output)])

    assert exit_code == 2
    assert "Unknown checklist table layout" in capsys.readouterr().err
    assert not output.exists()


def test_validator_rejects_incoherent_docling_pointer_index(tmp_path: Path) -> None:
    source = tmp_path / "input.docx"
    make_checklist(source)
    payload = digest_checklist(source)
    reference = payload["items"][0]["source_reference"]
    reference["docling_ref"] = "#/tables/99"

    with pytest.raises(ValueError, match="index does not match"):
        validate_digest(payload)
