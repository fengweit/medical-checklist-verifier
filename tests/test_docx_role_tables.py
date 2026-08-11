from pathlib import Path

from docx import Document

from medical_checklist_verifier.checklist_digester import digest_checklist


def test_bsi_role_table_treats_item_as_requirement_not_outcome(tmp_path: Path) -> None:
    source = tmp_path / "bsi.docx"
    document = Document()
    document.add_heading("BSI Completeness Checklist", level=1)
    table = document.add_table(rows=1, cols=4)
    headers = [
        "Section Title",
        "Item",
        "Location of the requested information",
        "BSI Completeness Check (To be completed by BSI)",
    ]
    for index, value in enumerate(headers):
        table.cell(0, index).text = value
    row = table.add_row().cells
    row[0].text = "Overview"
    row[1].text = "Cover letter"
    row[2].text = "Refer to approval sheet."
    row[3].text = "YES NO"
    document.save(str(source))

    result = digest_checklist(source)

    item = result["items"][0]
    assert item["label"] == "Cover letter"
    assert item["section_title"] == "Overview"
    assert item["declared_location"] == "Refer to approval sheet."
    assert [requirement["text"] for requirement in item["requirements"]] == [
        "Cover letter"
    ]
    assert item["requirements"][0]["source_reference"]["cell_index"] == 1
    assert "YES NO" not in str(item["requirements"])


def test_headerless_following_table_inherits_subject_schema(tmp_path: Path) -> None:
    source = tmp_path / "subject.docx"
    document = Document()
    document.add_heading("MDR Checklist", level=1)

    first = document.add_table(rows=2, cols=3)
    first.cell(0, 0).text = "Subject"
    first.cell(0, 1).text = "Reference to TD document (to be completed by client)"
    first.cell(0, 2).text = "Submission check (to be completed by client)"
    first.cell(1, 0).text = "1.1 (a) product or trade name and intended purpose;"
    first.cell(1, 1).text = "Document ID"
    first.cell(1, 2).text = "YES NO N/A"

    following = document.add_table(rows=1, cols=3)
    following.cell(0, 0).text = "1.1 (b) the Basic UDI-DI assigned by the manufacturer;"
    following.cell(0, 1).text = "Document ID"
    following.cell(0, 2).text = "YES NO N/A"
    document.save(str(source))

    result = digest_checklist(source)

    assert [item["label"] for item in result["items"]] == [
        "1.1 (a) product or trade name and intended purpose;",
        "1.1 (b) the Basic UDI-DI assigned by the manufacturer;",
    ]
    assert result["items"][1]["declared_location"] == "Document ID"
    assert result["items"][1]["source_reference"] == {
        "kind": "docling_table_cell",
        "docling_ref": "#/tables/1",
        "table_index": 1,
        "row_index": 0,
        "cell_index": 0,
    }


def test_docx_skips_merged_section_comments_and_conclusion_table(
    tmp_path: Path,
) -> None:
    source = tmp_path / "administrative.docx"
    document = Document()
    document.add_heading("Checklist", level=1)
    table = document.add_table(rows=1, cols=4)
    headers = ["Section Title", "Item", "Location", "Completeness Check"]
    for index, value in enumerate(headers):
        table.cell(0, index).text = value

    section = table.add_row().cells
    merged = section[0].merge(section[1])
    merged.text = "1. Device Description"

    item = table.add_row().cells
    item[0].text = "1. Device Description"
    item[1].text = "1.1 General description"
    item[2].text = "Document 1"
    item[3].text = "YES NO"

    comments = table.add_row().cells
    comments[0].text = "1. Device Description"
    comments[1].text = "BSI Comments"

    supplemental = document.add_table(rows=1, cols=3)
    supplemental.cell(0, 0).text = "Item"
    supplemental.cell(0, 1).text = "Location"
    supplemental.cell(0, 2).text = "Completeness Check"

    conclusion = document.add_table(rows=1, cols=3)
    conclusion.cell(0, 0).text = (
        "Required documentation received - Continue to formal "
        "Technical Documentation Review"
    )
    conclusion.cell(0, 1).text = (
        "Required documentation not received - Do not continue to formal "
        "Technical Documentation Review"
    )
    conclusion.cell(0, 2).text = "Technical Documentation Completeness check exempt"
    document.save(str(source))

    result = digest_checklist(source)

    assert [item["label"] for item in result["items"]] == ["1.1 General description"]
    assert {item["reason"] for item in result["exclusions"]} == {
        "administrative_or_reviewer_row",
        "conclusion_or_signoff_table",
        "section_divider_row",
    }
