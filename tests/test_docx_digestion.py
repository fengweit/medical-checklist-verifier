import json
from pathlib import Path

from docx import Document

from medical_checklist_verifier.checklist_digester import digest_checklist


def make_checklist(path: Path) -> None:
    document = Document()
    document.add_heading("Technical Documentation Checklist", level=1)
    table = document.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Requirements"
    table.rows[0].cells[2].text = "Reference"

    row = table.add_row().cells
    row[0].text = "1. Device description"
    row[1].text = "State the manufacturer name; State the intended purpose."
    row[2].text = "MDR Annex II §1.1"

    row = table.add_row().cells
    row[0].text = "2. Labelling"
    row[1].text = "Provide the label artwork."
    row[2].text = "MDR Annex II §2"
    document.save(path)


def test_docx_table_becomes_referenced_items_and_atomic_requirements(
    tmp_path: Path,
) -> None:
    source = tmp_path / "checklist.docx"
    make_checklist(source)

    result = digest_checklist(source)

    assert result["schema_version"] == "checklist-digester.v1"
    assert result["extraction"]["library"] == "docling"
    assert result["extraction"]["conversion_status"] == "success"
    assert result["extraction"]["interpreter"] == {
        "name": "docling-structural",
        "version": "3.0.0",
    }
    assert result["extraction"]["adapter_mode"] == "semantic_tables"
    assert result["extraction"]["docling_schema_name"] == "DoclingDocument"
    assert result["checklist"]["title"] == "Technical Documentation Checklist"
    assert result["checklist"]["source"]["filename"] == "checklist.docx"
    assert len(result["checklist"]["source"]["sha256"]) == 64

    items = result["items"]
    assert [item["label"] for item in items] == [
        "1. Device description",
        "2. Labelling",
    ]
    assert items[0]["declared_reference"] == "MDR Annex II §1.1"
    assert items[0]["source_reference"] == {
        "kind": "docling_table_cell",
        "docling_ref": "#/tables/0",
        "table_index": 0,
        "row_index": 1,
        "cell_index": 0,
    }
    assert [requirement["text"] for requirement in items[0]["requirements"]] == [
        "State the manufacturer name",
        "State the intended purpose.",
    ]
    assert items[0]["requirements"][1]["source_reference"] == {
        "kind": "docling_table_cell",
        "docling_ref": "#/tables/0",
        "table_index": 0,
        "row_index": 1,
        "cell_index": 1,
        "segment_index": 1,
    }
    assert items[0]["id"].startswith("item-")
    assert items[0]["requirements"][0]["id"].startswith(
        items[0]["id"] + "-requirement-"
    )

    # The output is directly serializable for later agent sessions.
    json.dumps(result, ensure_ascii=False)


def test_docx_digestion_is_deterministic(tmp_path: Path) -> None:
    source = tmp_path / "checklist.docx"
    make_checklist(source)

    assert digest_checklist(source) == digest_checklist(source)
